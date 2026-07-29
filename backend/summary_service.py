"""Generates the one-page SOP Change Summary .docx that accompanies every
tracked-changes SOP revision (per SOP_Change_Summary_Template.docx already
delivered for this program). Built with python-docx since this document has
no track-changes content -- it's a clean, generated summary page.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color


def generate_change_summary(output_path, sop_title, sop_number, site_name, prev_version, new_version,
                             revision_date, prepared_by, requirement, gap_description, change_heading,
                             change_paragraphs, ai_mock=False):
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    title = doc.add_paragraph()
    run = title.add_run("SOP Change Summary")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Regulatory-alignment revision — for reviewer / approver use alongside the tracked-changes SOP")
    sub_run.italic = True
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if ai_mock:
        warn = doc.add_paragraph()
        wr = warn.add_run("⚠ Generated in OFFLINE HEURISTIC MODE (no live AI connected) — treat all drafted text as placeholder pending a live-AI regeneration.")
        wr.bold = True
        wr.font.size = Pt(9)
        wr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    info_table = doc.add_table(rows=5, cols=4)
    info_table.style = "Table Grid"
    rows_data = [
        ("SOP Title", sop_title, "", ""),
        ("SOP Number", sop_number, "Site", site_name),
        ("Previous Version", prev_version, "New Version", new_version),
        ("Revision Date", revision_date, "Prepared By", prepared_by),
        ("Reason for Revision", "Close regulatory compliance gap identified during multi-site SOP / RTM review.", "", ""),
    ]
    for i, (l1, v1, l2, v2) in enumerate(rows_data):
        row = info_table.rows[i]
        _set_cell_text(row.cells[0], l1, bold=True)
        _shade_cell(row.cells[0], "D9E2F3")
        if l2:
            row.cells[0].merge(row.cells[0]) if False else None
            _set_cell_text(row.cells[1], v1)
            _set_cell_text(row.cells[2], l2, bold=True)
            _shade_cell(row.cells[2], "D9E2F3")
            _set_cell_text(row.cells[3], v2)
        else:
            merged = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
            _set_cell_text(merged, v1)

    doc.add_paragraph()
    h = doc.add_paragraph()
    hr = h.add_run("Change Made")
    hr.bold = True
    hr.font.size = Pt(14)
    hr.font.color.rgb = NAVY

    change_table = doc.add_table(rows=2, cols=3)
    change_table.style = "Table Grid"
    hdr = change_table.rows[0]
    for i, text in enumerate(["Section Added / Amended", "Change Description", "Regulatory Requirement Addressed"]):
        _set_cell_text(hdr.cells[i], text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(hdr.cells[i], "1F3864")
    body_row = change_table.rows[1]
    _set_cell_text(body_row.cells[0], change_heading)
    _set_cell_text(body_row.cells[1], " ".join(change_paragraphs)[:600])
    _set_cell_text(body_row.cells[2], f"{requirement['source']} — {requirement['clause']}: {requirement['requirement_text'][:200]}")

    doc.add_paragraph()
    h2 = doc.add_paragraph()
    h2r = h2.add_run("Gap Addressed")
    h2r.bold = True
    h2r.font.size = Pt(14)
    h2r.font.color.rgb = NAVY
    doc.add_paragraph(gap_description)

    doc.add_paragraph()
    h3 = doc.add_paragraph()
    h3r = h3.add_run("Review & Approval")
    h3r.bold = True
    h3r.font.size = Pt(14)
    h3r.font.color.rgb = NAVY

    approval_table = doc.add_table(rows=4, cols=4)
    approval_table.style = "Table Grid"
    hdr2 = approval_table.rows[0]
    for i, text in enumerate(["Role", "Name", "Signature / Decision", "Date"]):
        _set_cell_text(hdr2.cells[i], text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(hdr2.cells[i], "1F3864")
    for i, role in enumerate(["Prepared by", "Site QA Reviewer", "Approver"], start=1):
        _set_cell_text(approval_table.rows[i].cells[0], role, bold=True)
        _shade_cell(approval_table.rows[i].cells[0], "D9E2F3")

    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run(
        "Note: this summary accompanies the SOP file issued with Track Changes enabled. "
        "Accept/reject individual edits directly in Word; this page is not a substitute for reviewing the redline."
    )
    nr.italic = True
    nr.font.size = Pt(8)
    nr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(output_path)
    return output_path
