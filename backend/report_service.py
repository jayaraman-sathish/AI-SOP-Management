"""Generates downloadable Word reports for the two things this app previously
only ever showed as live, on-screen tables: a per-SOP compliance check (RTM
run) and a site-to-site SOP comparison. Both land in the Reports tab so a
user can come back and download them again later instead of the result only
existing as long as the page stays open.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
STATUS_COLORS = {
    "Covered": RGBColor(0x1E, 0x7B, 0x34),
    "Partially Covered": RGBColor(0x9A, 0x5B, 0x00),
    "Not Covered": RGBColor(0xC0, 0x00, 0x00),
    "SOP Missing": RGBColor(0x8B, 0x00, 0x00),
}


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _title_block(doc, title, subtitle):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    t = doc.add_paragraph()
    r = t.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY
    s = doc.add_paragraph()
    sr = s.add_run(subtitle)
    sr.italic = True
    sr.font.size = Pt(9)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()


def generate_compliance_report(output_path, sop, site, rtm_results, ai_mock_any=False,
                                general_issues=None, general_ai_mock=False, general_offline_note=None,
                                initiated_by=None, initiated_at=None, regulations_checked=None, report_id=None):
    """
    sop: dict with sop_number, title, sop_category
    site: dict with code, name
    rtm_results: list of {req_code, source, clause, requirement_text, coverage_status, rationale, sop_id}
        (the same shape returned by the RTM background job's "results" list)
    general_issues: list of {issue_type, location, current_text, proposed_correction, why_needed} --
        writing-quality/completeness findings, independent of any specific regulatory requirement.
    regulations_checked: list of regulatory framework/"source" names the user selected for this run
        (e.g. ["FDA 21 CFR 211", "USP <790>"]) -- the run only evaluates requirements from these,
        instead of blindly running the entire requirement library regardless of relevance.
    """
    general_issues = general_issues or []
    regulations_checked = regulations_checked or []
    doc = Document()
    _title_block(
        doc,
        "SOP Compliance Report",
        f"{sop['sop_number']} — {sop['title']}  |  {site['code']} — {site['name']}",
    )
    if report_id is not None:
        idp = doc.add_paragraph()
        idr = idp.add_run(f"Report ID: RPT-{report_id}  (look this up in the app's Reports tab or Audit Trail to verify this document and confirm when it was generated)")
        idr.bold = True
        idr.font.size = Pt(9)
        idr.font.color.rgb = NAVY
    if initiated_by or initiated_at:
        meta = doc.add_paragraph()
        mr = meta.add_run(f"Initiated by {initiated_by or 'unknown'} on {initiated_at or 'unknown'}")
        mr.font.size = Pt(9)
        mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if regulations_checked:
        regs = doc.add_paragraph()
        rr = regs.add_run(f"Regulations checked in this run: {', '.join(regulations_checked)}")
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if ai_mock_any:
        warn = doc.add_paragraph()
        wr = warn.add_run("⚠ One or more assessments below were generated in OFFLINE HEURISTIC MODE (no live AI connected) — treat those as placeholder pending a live-AI re-run.")
        wr.bold = True
        wr.font.size = Pt(9)
        wr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        doc.add_paragraph()

    counts = {}
    for r in rtm_results:
        counts[r["coverage_status"]] = counts.get(r["coverage_status"], 0) + 1
    total = len(rtm_results)
    checked = total - counts.get("SOP Missing", 0)

    scope_note = (
        f"across the {len(regulations_checked)} selected regulation(s)" if regulations_checked
        else "across the full requirement library"
    )
    summary = doc.add_paragraph()
    sr = summary.add_run(
        f"{total} regulatory requirements evaluated {scope_note}. {checked} were within this SOP's actual scope "
        f"(the rest are correctly not applicable -- this SOP wasn't relevant to their subject matter). "
        f"Of those in scope: {counts.get('Covered', 0)} Covered, {counts.get('Partially Covered', 0)} "
        f"Partially Covered, {counts.get('Not Covered', 0)} Not Covered."
    )
    sr.font.size = Pt(10)
    doc.add_paragraph()

    h = doc.add_paragraph()
    hr = h.add_run("Compliant Points")
    hr.bold = True
    hr.font.size = Pt(14)
    hr.font.color.rgb = NAVY
    covered = [r for r in rtm_results if r["coverage_status"] == "Covered"]
    _result_table(doc, covered, "No fully-covered requirements were found in scope for this SOP.")

    doc.add_paragraph()
    h2 = doc.add_paragraph()
    h2r = h2.add_run("Non-Compliant / Partially Compliant Points")
    h2r.bold = True
    h2r.font.size = Pt(14)
    h2r.font.color.rgb = NAVY
    non_compliant = [r for r in rtm_results if r["coverage_status"] in ("Partially Covered", "Not Covered")]
    _result_table(doc, non_compliant, "No gaps were found in requirements within this SOP's scope.")

    missing = [r for r in rtm_results if r["coverage_status"] == "SOP Missing"]
    if missing:
        doc.add_paragraph()
        h3 = doc.add_paragraph()
        h3r = h3.add_run("Not Applicable to This SOP")
        h3r.bold = True
        h3r.font.size = Pt(14)
        h3r.font.color.rgb = NAVY
        note = doc.add_paragraph()
        nr = note.add_run(
            "These requirements were checked but don't pertain to this SOP's subject matter -- listed here for "
            "completeness, not as findings against this document."
        )
        nr.italic = True
        nr.font.size = Pt(9)
        nr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        _result_table(doc, missing, "", compact=True)

    doc.add_paragraph()
    h4 = doc.add_paragraph()
    h4r = h4.add_run("General — Writing Quality & Completeness")
    h4r.bold = True
    h4r.font.size = Pt(14)
    h4r.font.color.rgb = NAVY
    note4 = doc.add_paragraph()
    n4r = note4.add_run(
        "Grammar, spelling, sentence-formation, and missing-form/template issues found in this SOP's own "
        "text -- independent of regulatory compliance, which is covered in the sections above."
    )
    n4r.italic = True
    n4r.font.size = Pt(9)
    n4r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    if general_ai_mock:
        gwarn = doc.add_paragraph()
        gwr = gwarn.add_run(f"⚠ {general_offline_note or 'General review requires a live Claude API key.'}")
        gwr.bold = True
        gwr.font.size = Pt(9)
        gwr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    _general_issues_table(doc, general_issues, "No writing-quality or completeness issues were found.")

    doc.save(output_path)
    return output_path


def _general_issues_table(doc, issues, empty_note):
    if not issues:
        p = doc.add_paragraph(empty_note)
        if p.runs:
            p.runs[0].italic = True
        return
    cols = ["Issue Type", "Location", "Current Text", "Proposed Correction", "Why Needed"]
    table = doc.add_table(rows=1 + len(issues), cols=len(cols))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, text in enumerate(cols):
        _set_cell_text(hdr.cells[i], text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9)
        _shade_cell(hdr.cells[i], "1F3864")
    for ri, issue in enumerate(issues, start=1):
        row = table.rows[ri]
        _set_cell_text(row.cells[0], issue.get("issue_type", ""), bold=True, size=9)
        _set_cell_text(row.cells[1], issue.get("location", ""), size=9)
        _set_cell_text(row.cells[2], (issue.get("current_text") or "")[:300], size=9)
        _set_cell_text(row.cells[3], (issue.get("proposed_correction") or "")[:300], size=9)
        _set_cell_text(row.cells[4], (issue.get("why_needed") or "")[:300], size=9)


def generate_compliance_summary_report(output_path, sop, site, rtm_results, general_issues=None,
                                        initiated_by=None, initiated_at=None, regulations_checked=None, report_id=None):
    """
    A short, one-page companion to the full Compliance Report -- who ran it,
    when, and the headline numbers, without the detailed per-requirement
    tables. Mirrors the existing Redline + Change Summary pairing so a
    reviewer who just wants the top line doesn't have to open the full report.
    """
    general_issues = general_issues or []
    regulations_checked = regulations_checked or []
    doc = Document()
    _title_block(
        doc,
        "SOP Compliance Summary",
        f"{sop['sop_number']} — {sop['title']}  |  {site['code']} — {site['name']}",
    )
    if report_id is not None:
        idp = doc.add_paragraph()
        idr = idp.add_run(f"Report ID: RPT-{report_id}  (look this up in the app's Reports tab or Audit Trail to verify this document and confirm when it was generated)")
        idr.bold = True
        idr.font.size = Pt(9)
        idr.font.color.rgb = NAVY
    meta = doc.add_paragraph()
    mr = meta.add_run(f"Initiated by {initiated_by or 'unknown'} on {initiated_at or 'unknown'}")
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if regulations_checked:
        regs = doc.add_paragraph()
        rr = regs.add_run(f"Regulations checked in this run: {', '.join(regulations_checked)}")
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    counts = {}
    for r in rtm_results:
        counts[r["coverage_status"]] = counts.get(r["coverage_status"], 0) + 1
    total = len(rtm_results)
    checked = total - counts.get("SOP Missing", 0)

    cols = ["Metric", "Count"]
    scope_label = (
        f"Regulatory requirements evaluated ({len(regulations_checked)} regulation(s) selected)"
        if regulations_checked else "Regulatory requirements evaluated (full library)"
    )
    stat_rows = [
        (scope_label, total),
        ("Within this SOP's actual scope", checked),
        ("Compliant", counts.get("Covered", 0)),
        ("Partially Compliant", counts.get("Partially Covered", 0)),
        ("Non-Compliant", counts.get("Not Covered", 0)),
        ("Not Applicable to this SOP", counts.get("SOP Missing", 0)),
        ("General writing/completeness issues", len(general_issues)),
    ]
    table = doc.add_table(rows=1 + len(stat_rows), cols=len(cols))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, text in enumerate(cols):
        _set_cell_text(hdr.cells[i], text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(hdr.cells[i], "1F3864")
    for ri, (label, value) in enumerate(stat_rows, start=1):
        row = table.rows[ri]
        _set_cell_text(row.cells[0], label, size=10)
        _set_cell_text(row.cells[1], value, bold=True, size=10)

    non_compliant = [r for r in rtm_results if r["coverage_status"] in ("Partially Covered", "Not Covered")]
    if non_compliant:
        doc.add_paragraph()
        h = doc.add_paragraph()
        hr = h.add_run("Top Findings Needing Attention")
        hr.bold = True
        hr.font.size = Pt(13)
        hr.font.color.rgb = NAVY
        for r in non_compliant[:5]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{r.get('req_code','')} ({r.get('source','')} {r.get('clause','')}) — {r.get('coverage_status','')}")
            run.bold = True
            run.font.size = Pt(10)
        if len(non_compliant) > 5:
            more = doc.add_paragraph()
            more.add_run(f"...and {len(non_compliant) - 5} more. See the full Compliance Report for details.").italic = True

    doc.save(output_path)
    return output_path


def _result_table(doc, results, empty_note, compact=False):
    if not results:
        p = doc.add_paragraph(empty_note)
        p.runs[0].italic = True if p.runs else None
        return
    cols = ["Req Code", "Source / Clause", "Status", "Rationale"] if not compact else ["Req Code", "Source / Clause", "Requirement"]
    table = doc.add_table(rows=1 + len(results), cols=len(cols))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, text in enumerate(cols):
        _set_cell_text(hdr.cells[i], text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(hdr.cells[i], "1F3864")
    for ri, r in enumerate(results, start=1):
        row = table.rows[ri]
        _set_cell_text(row.cells[0], r.get("req_code", ""), bold=True, size=9)
        _set_cell_text(row.cells[1], f"{r.get('source','')} {r.get('clause','')}".strip(), size=9)
        if not compact:
            status_color = STATUS_COLORS.get(r.get("coverage_status"), None)
            _set_cell_text(row.cells[2], r.get("coverage_status", ""), bold=True, color=status_color, size=9)
            _set_cell_text(row.cells[3], (r.get("rationale") or "")[:500], size=9)
        else:
            _set_cell_text(row.cells[2], (r.get("requirement_text") or "")[:200], size=9)


def generate_comparison_report(output_path, sops_compared, findings, ai_mock=False, report_id=None):
    """
    sops_compared: list of {site_code, site_name, sop_number, title, version_label}
    findings: list of {process_step, site_values (dict), classification, note}

    version_label is printed alongside each SOP so a re-comparison run after
    an SOP revision produces a report that's unambiguously distinguishable
    from an earlier comparison of the same SOP pair -- not just by timestamp.
    """
    doc = Document()
    scope = "; ".join(
        f"{s['site_code']} — {s['sop_number']} ({s.get('version_label') or 'version unknown'}) {s['title']}"
        for s in sops_compared
    )
    _title_block(doc, "SOP Site Comparison Report", scope)

    if report_id is not None:
        idp = doc.add_paragraph()
        idr = idp.add_run(f"Report ID: RPT-{report_id}  (look this up in the app's Reports tab or Audit Trail to verify this document and confirm when it was generated)")
        idr.bold = True
        idr.font.size = Pt(9)
        idr.font.color.rgb = NAVY

    if ai_mock:
        warn = doc.add_paragraph()
        wr = warn.add_run("⚠ Generated in OFFLINE HEURISTIC MODE (no live AI connected) — treat all findings as placeholder pending a live-AI re-run.")
        wr.bold = True
        wr.font.size = Pt(9)
        wr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        doc.add_paragraph()

    summary = doc.add_paragraph()
    sr = summary.add_run(f"{len(findings)} concrete difference(s) identified across {len(sops_compared)} SOP(s).")
    sr.font.size = Pt(10)
    doc.add_paragraph()

    if not findings:
        doc.add_paragraph("No differences were identified -- the compared SOPs appear consistent on the points reviewed.")
    else:
        site_codes = sorted({s["site_code"] for s in sops_compared})
        cols = ["Process Step"] + site_codes + ["Classification", "Note", "Recommended Action"]
        table = doc.add_table(rows=1 + len(findings), cols=len(cols))
        table.style = "Table Grid"
        hdr = table.rows[0]
        for i, text in enumerate(cols):
            _set_cell_text(hdr.cells[i], text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9)
            _shade_cell(hdr.cells[i], "1F3864")
        for ri, f in enumerate(findings, start=1):
            row = table.rows[ri]
            _set_cell_text(row.cells[0], f.get("process_step", ""), bold=True, size=9)
            site_values = f.get("site_values", {}) or {}
            for ci, code in enumerate(site_codes, start=1):
                _set_cell_text(row.cells[ci], site_values.get(code, ""), size=9)
            _set_cell_text(row.cells[len(site_codes) + 1], f.get("classification", ""), size=9)
            _set_cell_text(row.cells[len(site_codes) + 2], f.get("note", ""), size=9)
            _set_cell_text(row.cells[len(site_codes) + 3], f.get("recommended_action", ""), size=9)

    doc.save(output_path)
    return output_path
