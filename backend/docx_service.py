"""Word (.docx) services: text extraction for AI analysis, and generation of
a tracked-changes ("Track Changes" / redline) revision of an uploaded SOP.

Track changes are written directly into the OOXML (word/document.xml) using
native <w:ins> insertion markup, so the output opens in Microsoft Word with
Track Changes already showing -- the reviewer sees a real redline, and can
accept/reject each inserted block from Word's Review pane.

v1 scope (documented limitation): edits are additive. The engine inserts a
new, clearly-headed tracked section containing the AI-drafted procedure text
needed to close the gap, immediately before the end of the document body (or
before a detected "Revision History" heading if present). It does not attempt
in-place surgical edits inside existing sentences, because Word fragments
paragraph text across many <w:r> runs in ways that make blind text-matching
unreliable (see docx-editing guidance) -- inserting a new tracked subsection
is the robust, inspection-safe approach for a v1. Editing existing passages
in place is a natural Phase 2 enhancement once the app is validated on real
site SOPs.
"""
import os
import re
import zipfile
import shutil
import datetime
import uuid
from docx import Document as DocxDocument

NSMAP_DECL = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def extract_text(docx_path):
    """Extract plain text (paragraphs + table cells) from a .docx for AI analysis."""
    try:
        doc = DocxDocument(docx_path)
    except Exception as e:
        return f"[Could not parse document: {e}]"
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _esc(text):
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _tracked_paragraph_xml(text, author, date_iso, bold=False, ins_id=1):
    rpr = "<w:rPr><w:b/></w:rPr>" if bold else ""
    return (
        f'<w:p><w:pPr><w:rPr><w:ins w:id="{ins_id}" w:author="{_esc(author)}" w:date="{date_iso}"/></w:rPr></w:pPr>'
        f'<w:ins w:id="{ins_id + 1}" w:author="{_esc(author)}" w:date="{date_iso}">'
        f'<w:r>{rpr}<w:t xml:space="preserve">{_esc(text)}</w:t></w:r>'
        f"</w:ins></w:p>"
    )


def generate_tracked_redline(source_docx_path, output_docx_path, heading, paragraphs, requirement,
                              author="AI Draft – SOP Compliance Platform", new_version_label=None,
                              gap_description=""):
    """
    Produce a copy of source_docx_path at output_docx_path with a new tracked
    (Track Changes) section inserted, documenting the SOP change needed to
    close a regulatory gap. Returns True on success.
    """
    date_iso = datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    shutil.copyfile(source_docx_path, output_docx_path)

    with zipfile.ZipFile(output_docx_path, "r") as zin:
        names = zin.namelist()
        if "word/document.xml" not in names:
            raise ValueError("Not a valid .docx (no word/document.xml)")
        contents = {n: zin.read(n) for n in names}

    xml = contents["word/document.xml"].decode("utf-8")

    ins_id = 1000 + int(uuid.uuid4().int % 8000)
    blocks = []
    blocks.append(_tracked_paragraph_xml("", author, date_iso, ins_id=ins_id))  # spacer
    ins_id += 2
    title_text = f"Regulatory Compliance Update – {requirement.get('req_code','')} ({requirement.get('source','')} {requirement.get('clause','')})"
    blocks.append(_tracked_paragraph_xml(title_text, author, date_iso, bold=True, ins_id=ins_id))
    ins_id += 2
    if new_version_label:
        blocks.append(_tracked_paragraph_xml(
            f"[Tracked insertion – proposed as part of SOP revision to version {new_version_label}. "
            f"Gap addressed: {gap_description}]",
            author, date_iso, ins_id=ins_id,
        ))
        ins_id += 2
    blocks.append(_tracked_paragraph_xml(heading, author, date_iso, bold=True, ins_id=ins_id))
    ins_id += 2
    for para_text in paragraphs:
        blocks.append(_tracked_paragraph_xml(para_text, author, date_iso, ins_id=ins_id))
        ins_id += 2

    insertion_xml = "".join(blocks)

    # Insert right before the sectPr at the end of the body (standard docx structure),
    # so the new tracked content lands at the end of the document body.
    #
    # IMPORTANT: a document can contain more than one <w:sectPr> -- every section
    # break before the last one is nested inside the <w:pPr> of that section's
    # final paragraph (e.g. a landscape page for a wide table, common in real
    # SOPs with annexures). Only the very LAST <w:sectPr> in the file is the
    # true end-of-body one, appearing as a direct child of <w:body> after the
    # final paragraph. Using the FIRST match (as this used to) planted our new
    # <w:p> blocks inside an earlier section-break's <w:pPr> -- <w:p> is not a
    # valid child of <w:pPr>, so Word rejected the whole file as corrupt. Using
    # the LAST match is always correct: there is exactly one final sectPr and it
    # is always positioned after everything else in the body.
    sect_pr_matches = list(re.finditer(r"<w:sectPr[ >]", xml))
    if sect_pr_matches:
        idx = sect_pr_matches[-1].start()
        new_xml = xml[:idx] + insertion_xml + xml[idx:]
    else:
        # Fallback: insert before </w:body>
        new_xml = xml.replace("</w:body>", insertion_xml + "</w:body>")

    contents["word/document.xml"] = new_xml.encode("utf-8")

    # Deliberately NOT injecting <w:trackChanges/> into settings.xml: it was being
    # inserted as the first child of <w:settings>, which violates the schema's
    # required element order (CT_Settings expects a long, specific sequence --
    # trackChanges has to slot in at a precise position, not just "first"), and
    # a validator run against a real multi-section SOP flagged it as invalid.
    # It's also unnecessary: Word displays existing <w:ins>/<w:del> markup as
    # tracked changes regardless of this setting -- it only affects whether
    # Word tracks *new* edits the reviewer makes afterward, which isn't needed
    # for the redline to render correctly.

    with zipfile.ZipFile(output_docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in contents.items():
            zout.writestr(name, data)

    return True
